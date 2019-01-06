import re
import json
import sys
from docx import Document

# Method check if string has Yes/No
def checkYesNo(str):
    if (str == "Yes" or str == "yes" or str == "No" or str == "no"):
        return True
    return False

# Method return description and question from string
def fiStepDescriptionQuestion(str):
    objDesQue = {}
    arrDesQue = str.split('.')

    if('?' not in str):
        objDesQue['fiStpDsc'] = str
    else:
        newDes = ""
        for description in arrDesQue[:-1]:
            newDes += description+"."
        objDesQue['fiStpDsc'] = newDes
        objDesQue['fiStpQst'] = arrDesQue[-1]

    return objDesQue

############# Main program ###############

removeCharacters = ['\n','\t']
IsFaultIsolation = False
FI_Dict = {}
FI_Label = True
FI_Descriptoin = ""
FI_Num = 1


document = Document('DOC-Troubleshooting_table.docx')
for para in document.paragraphs:
    if (FI_Label == True):
        FI_Dict['lbl'] = para.text
        FI_Label = False
    else:
        FI_Descriptoin+=para.text + "\n"

FI_Dict['des'] = FI_Descriptoin
FI_Dict['PG'] = []

tables = document.tables
for table in tables:
    for row in table.rows:
        FI_row = []
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if not checkYesNo(paragraph.text):
                    # print(paragraph.text)
                    FI_row.append(paragraph.text)
                    if len(FI_row) == 3:
                        newNumberObj = {}
                        newNumberObj['n'] = FI_Num
                        newNumberObj['docObj'] = fiStepDescriptionQuestion(FI_row[0])
                        newNumberObj['Y'] = FI_row[1]
                        newNumberObj['N'] = FI_row[2]
                        FI_Dict['PG'].append(newNumberObj)
                        FI_Num+=1

#print(FI_Dict)

print(json.dumps(FI_Dict, indent=4, sort_keys=True))

# print("-" + re.sub('\n','',FI_Dict['table'][1]['yes']) + "-")