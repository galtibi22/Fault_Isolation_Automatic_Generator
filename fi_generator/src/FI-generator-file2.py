import json
import re
import sys
import argparse
from docx import Document

# Method return to Step description and question from string
def fiStepDescriptionQuestion(str):
    htmlDataObj = {}
    htmlDataObjAraay = []

    arrDesQue = str.split('.')

    if('?' not in str):
        objDes = {'htmlType':'fiStpDsc', 'txt' : str}
        htmlDataObjAraay.append(objDes)
    else:
        newDes = ""
        for description in arrDesQue[:-1]:
            newDes += description+"."
        objDes = {'htmlType':'fiStpDsc', 'txt' : newDes}
        htmlDataObjAraay.append(objDes)
        objQue = {'fiStpQst':'fiStpDsc', 'txt' : arrDesQue[-1]}
        htmlDataObjAraay.append(objQue)

    htmlDataObj['htmlData'] = htmlDataObjAraay
    return htmlDataObj


# Method return description about FI flow
def fiMainDescription(Header_Name, Header_Description):

    FI_Descriptoin = ""
    FI_Main_HTML_Data = []
    FI_Main_HTML_Data_Obj = {}
    FI_Main_HTML_Obj = {}

    for i in range(1,4):
        FI_Descriptoin += Header_Name[i] + ": " + Header_Description[i] + "\n"

    FI_Main_HTML_Data_Obj['htmlType'] = 'fiTitle'
    FI_Main_HTML_Data_Obj['txt'] = Header_Name[0] + ": " + Header_Description[0]
    FI_Main_HTML_Data.append(FI_Main_HTML_Data_Obj)

    FI_Main_HTML_Data_Obj = {}
    FI_Main_HTML_Data_Obj['htmlType'] = 'fiStpDes'
    FI_Main_HTML_Data_Obj['txt'] = FI_Descriptoin
    FI_Main_HTML_Data.append(FI_Main_HTML_Data_Obj)

    FI_Main_HTML_Obj['htmlData'] = FI_Main_HTML_Data

    return FI_Main_HTML_Obj


# Method return remove characters
def removeCharsFunc(str):
    for rchar in removeCharacters:
        str = str.replace(rchar,'')
    return removeRegexFunc(str)

# Method return remove start characters with regex
def removeRegexFunc(str):
    return re.sub(r'(^ |^[0-9]{2}[A-Z])', '', str)


# Method return Task Object to 'Y' (Yes option only)
def fiTaskYes2(str, yesOption, noOption):
    yesObj = {}

    yesObj['to'] = removeCharsFunc(str)
    yesObj['rtN'] = removeCharsFunc(noOption)
    yesObj['msgRt'] = "1"
    yesObj['typ'] = "1"
    yesObj['msgRtIx'] = "0"
    yesObj['tskNm'] = removeCharsFunc(str)
    yesObj['rtY'] = removeCharsFunc(yesOption)

    return yesObj


#############################################################################################
#############################################################################################
###################################                    ######################################
###################################    Main program    ######################################
###################################                    ######################################
#############################################################################################
#############################################################################################

removeCharacters = ['\n','\t','\u200e']
FI_Label = True
FI_Num = 0

FI_Txt_Header = {}
FI_Txt_Header_Exist = False

# parser = argparse.ArgumentParser()
# parser.add_argument("file")
# args = parser.parse_args()
# document = Document(args.file)

document = Document('C:\\Users\\eden.SPIDERSERVICES\\Desktop\\docx\\DOC-Troubleshooting_table_testing.docx')

FI_Array_List = []
FI_Array = []

tables = document.tables
print(len(tables[0].rows[0].cells))

for table in tables:
    for row in table.rows:
        FI_row = []
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                FI_row.append(paragraph.text)

                if(len(FI_row) == 8):

                    if not FI_Txt_Header_Exist:
                        for i in range(0,8):
                            FI_Txt_Header[i] = FI_row[i]
                        FI_Txt_Header_Exist = True

                    else:
                        # Main text info to object (Failure No., Tested Unit, Severity, Platform)
                        FI_Num=0
                        newNumberObj = {}
                        newNumberObj['n'] = str(FI_Num)
                        FI_Num+=1
                        newNumberObj['htmlObj'] = fiMainDescription(FI_Txt_Header, FI_row)
                        FI_Array.append(newNumberObj)
                        Number_Of_Actions = 0

                        for i in range(4, 8):
                            if FI_row[i] != "":
                                Number_Of_Actions +=1

                        for i in range(4, 8):
                            if FI_row[i] != "":
                                newNumberObj = {}
                                newNumberObj['n'] = str(FI_Num)
                                if(i == 4):
                                    # Step
                                    newNumberObj['Y'] = {'to' : str(Number_Of_Actions+1), 'typ' : '0'}
                                    newNumberObj['N'] = {'to' : str(FI_Num+1), 'typ' : '0'}
                                else:
                                    # Task
                                    newNumberObj['Y'] = fiTaskYes2(FI_row[i], str(Number_Of_Actions+1), str(FI_Num+1))
                                    newNumberObj['N'] = { 'typ' : '4' }
                                newNumberObj['htmlObj'] = fiStepDescriptionQuestion(FI_Txt_Header[i] + ": " + FI_row[i])
                                FI_Array.append(newNumberObj)
                                FI_Num+=1


                        ###### Default end FI numbers ######
                        FI_Array.append({"n": str(FI_Num), "htmlObj": {"htmlData": [{"htmlType": "fiNegEnd"}]},
                                         "N": {"typ": "4"}, "Y": {"typ": "4"}})
                        FI_Num += 1
                        FI_Array.append({"n": str(FI_Num), "htmlObj": {"htmlData": [{"htmlType": "fiPosEnd"}]},
                                         "N": {"typ": "4"}, "Y": {"typ": "4"}})
                        FI_Array_List.append(FI_Array)
                        FI_Array = []



print(json.dumps(FI_Array_List, indent=4, sort_keys=True))