from glob import glob
import json
import re
import sys
import argparse
import os
import requests
import subprocess
from docx import Document
import time
#import win32com.client as win32
#from win32com.client import constants

# Method check if string has Yes/No

def checkYesNo(str):
    if (str == "Yes" or str == "yes" or str == "No" or str == "no"):
        return True
    return False

# Method return remove characters
def removeCharsFunc(str):
    for rchar in removeCharacters:
        str = str.replace(rchar,'')
    return removeRegexFunc(str)

# Method return remove start characters with regex
def removeRegexFunc(str):
    return re.sub(r'(^ |^[0-9]{2}[A-Z])', '', str)


# Method return to Step description and question from string
def fiStepDescriptionQuestion(str):
    htmlDataObj = {}
    htmlDataObjAraay = []

    arrDesQue = str.split('.')

    if('?' not in str):
        objDes = {'htmlType':'fiStpPrc', 'txt' : removeCharsFunc(str)}
        htmlDataObjAraay.append(objDes)
    else:
        newDes = ""
        for description in arrDesQue[:-1]:
            newDes += description+"."
        objDes = {'htmlType':'fiStpPrc', 'txt' : removeCharsFunc(newDes)}
        htmlDataObjAraay.append(objDes)
        objQue = {'htmlType':'fiStpQst', 'txt' : removeCharsFunc(arrDesQue[-1])}
        htmlDataObjAraay.append(objQue)

    htmlDataObj['htmlData'] = htmlDataObjAraay
    return htmlDataObj



# Method return to Task description and question from string
def fiTaskDescriptionQuestion(str):
    htmlDataObj = {}
    htmlDataObjAraay = []

    arrDesQue = str.split('.')

    if('?' not in str):
        objDes = {'htmlType':'fiStpPrc', 'txt' : removeCharsFunc(str)}
        htmlDataObjAraay.append(objDes)
    else:
        newDes = ""
        for description in arrDesQue[:-1]:
            newDes += description+"."
        objDes = {'htmlType':'fiStpPrc', 'txt' : removeCharsFunc(newDes)}
        htmlDataObjAraay.append(objDes)

    htmlDataObj['htmlData'] = htmlDataObjAraay
    return htmlDataObj



# Method return Task Object to 'Y' (Yes option only)
def fiTaskYes(str, yesOption, noOption):
    yesObj = {}
    arrDesQue = str.split('.')
    newDes = ""

    if('?' not in str):
        newDes = str
    else:
        yesObj['msgRtIx'] = removeCharsFunc(arrDesQue[-1])
        for description in arrDesQue[:-1]:
            newDes += description+"."

    yesObj['to'] = removeCharsFunc(newDes[newDes.find("(")+1:newDes.find(")")])
    yesObj['rtN'] = removeCharsFunc(noOption)
    yesObj['msgRt'] = "1"
    yesObj['typ'] = "1"
    yesObj['tskNm'] = removeCharsFunc(newDes[:newDes.find("(")])
    yesObj['rtY'] = removeCharsFunc(yesOption)

    return yesObj


def save_as_docx_win(path):
        # Opening MS Word
        word = win32.gencache.EnsureDispatch('Word.Application')
        doc = word.Documents.Open(path)
        doc.Activate ()

        # Rename path with .docx
        new_file_abs = os.path.abspath(path)
        new_file_abs = re.sub(r'\.\w+$', '.docx', new_file_abs)

        # Save and Close
        word.ActiveDocument.SaveAs(
            new_file_abs, FileFormat=constants.wdFormatXMLDocument
        )
        doc.Close(False)

        return new_file_abs


def save_as_docx_mac(path):
    #print("path",path)
    #SOFFICE_PATH='../../Office.app/Contents/MacOS/soffice'
    SOFFICE_PATH="C:\Program Files\LibreOffice\program\soffice.exe"
    subprocess.call([SOFFICE_PATH, '--headless', '--convert-to', 'docx', path])
    f = open(path)
    return os.path.basename(f.name).replace("doc","docx")



def post_api_server(uri, jsondata):
    # defining the api-endpoint
    API_ENDPOINT = "http://127.0.0.1/api/figenerator/new/"

    headers = {
        'content-type' : "application/json"
    }
    r = requests.post(API_ENDPOINT+uri, data = jsondata, headers = headers)




#############################################################################################
#############################################################################################
###################################                    ######################################
###################################    Main program    ######################################
###################################                    ######################################
#############################################################################################
#############################################################################################

removeCharacters = ['\n','\t','\u200e']
FI_Label = True
FI_Descriptoin = ""
FI_Num = 0

parser = argparse.ArgumentParser()
parser.add_argument("source")
parser.add_argument("result")
args = parser.parse_args()
path = glob(args.source, recursive=True)
# document = Document(args.source)


# path = glob('C:\\Users\\eden.SPIDERSERVICES\\Desktop\\docx\\test\\DOC-Left-Wagon.doc', recursive=True)


if(path[0].endswith('docx')):
    print([path[0]])
    document = Document(path[0])
else:
    if (path[0].endswith('doc')):
       # if (os.name == 'posix'):
            print("use save_as_docx_mac")
            pathUrl=[save_as_docx_mac(path[0])]
            f=open(pathUrl[0])
            f.close()
            print("pathUrl retrun from save_as_docx is ",pathUrl)
            document = Document(pathUrl)
        #else:
         #   print("use save_as_docx_win")
          #  document = Document(save_as_docx_win(path[0]))


FI_Array = []

# Main text info to object
FI_Main_Info = {}
FI_Main_Info['n'] = str(FI_Num)
FI_Num+=1
FI_Main_HTML_Obj ={}
FI_Main_HTML_Data = []
FI_Main_HTML_Data_Obj = {}

for para in document.paragraphs:
    if (FI_Label == True):
        FI_Main_HTML_Data_Obj['htmlType'] = 'fiTitle'
        FI_Main_HTML_Data_Obj['txt'] = para.text
        FI_Label = False
        FI_Main_HTML_Data.append(FI_Main_HTML_Data_Obj)
        FI_Main_HTML_Data_Obj = {}
    else:
        FI_Descriptoin+=para.text + "\n"

FI_Main_HTML_Data_Obj['htmlType'] = 'fiStpDsc'
FI_Main_HTML_Data_Obj['txt'] = FI_Descriptoin
FI_Main_HTML_Data.append(FI_Main_HTML_Data_Obj)

FI_Main_HTML_Obj['htmlData'] = FI_Main_HTML_Data
FI_Main_Info['htmlObj'] = FI_Main_HTML_Obj
FI_Array.append(FI_Main_Info)

tables = document.tables

# print(len(tables[0].rows[0].cells))
for table in tables:
    for row in table.rows:
        FI_row = []
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if not checkYesNo(paragraph.text):
                    # print(paragraph.text)
                    FI_row.append(paragraph.text)
                    if len(FI_row) == 3:
                        des=FI_row[0]
                        newNumberObj = {}
                        newNumberObj['n'] = str(FI_Num)
                        if(('(' not in  FI_row[0]) and (')' not in  FI_row[0])):
                            # Step
                            newNumberObj['Y'] = {'to' : removeCharsFunc(FI_row[1]), 'typ' : '0'}
                            newNumberObj['N'] = {'to' : removeCharsFunc(FI_row[2]), 'typ' : '0'}
                            newNumberObj['htmlObj'] = fiStepDescriptionQuestion(FI_row[0])
                        else:
                            # Task
                            newNumberObj['Y'] = fiTaskYes(FI_row[0], FI_row[1], FI_row[2])
                            newNumberObj['N'] = { 'typ' : '4' }
                            newNumberObj['htmlObj'] = fiTaskDescriptionQuestion(FI_row[0])
                        FI_Array.append(newNumberObj)
                        FI_Num+=1


###### Default end FI numbers ######
FI_Array.append({ "n": str(FI_Num), "htmlObj": { "htmlData": [ { "htmlType": "fiNegEnd" } ] }, "N": { "typ": "4" }, "Y": { "typ": "4" } })
FI_Num+=1
FI_Array.append({ "n": str(FI_Num), "htmlObj": { "htmlData": [ { "htmlType": "fiPosEnd" } ] }, "N": { "typ": "4" }, "Y": { "typ": "4" } })


f= open(args.result,"w+")
jsonrequest = "{\"PG\":"+json.dumps(FI_Array, indent=4, sort_keys=True)+"}"
f.write(jsonrequest)
print(jsonrequest)

post_api_server(glob(args.result, recursive=True), jsonrequest)

