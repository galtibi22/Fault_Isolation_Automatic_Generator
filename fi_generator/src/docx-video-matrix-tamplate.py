import re
import json
import sys
from docx import Document

def checkYesNo(str):
    if (str == "Yes" or str == "yes" or str == "No" or str == "no"):
        return True
    return False

removeCharacters = ['\n','\t']
IsFaultIsolation = False
FI_Dict = {}
FI_Paragraphs = []
FI_Num = 1

document = Document('DOC-Video-Matrix.docx')
for para in document.paragraphs:
    print(para.text)
    FI_Paragraphs.append(para.text)

FI_Dict['paragraphs'] = FI_Paragraphs
FI_Dict['table'] = {}

tables = document.tables
for table in tables:
    for row in table.rows:
        FI_row = []
        for cell in row.cells:
            for paragraph in cell.paragraphs:

                if checkYesNo(paragraph.text):
                    IsFaultIsolation = True

                if (IsFaultIsolation == True):
                    if not checkYesNo(paragraph.text):
                        print(paragraph.text)
                        FI_row.append(paragraph.text)
                        if len(FI_row) == 3:
                            FI_Dict['table'][FI_Num] = {}
                            FI_Dict['table'][FI_Num]['description'] = FI_row[0]
                            FI_Dict['table'][FI_Num]['yes'] = FI_row[1]
                            FI_Dict['table'][FI_Num]['no'] = FI_row[2]
                            FI_Num+=1

print(FI_Dict)

print(json.dumps(FI_Dict, indent=4, sort_keys=True))

print("-" + re.sub('\n','',FI_Dict['table'][1]['yes']) + "-")