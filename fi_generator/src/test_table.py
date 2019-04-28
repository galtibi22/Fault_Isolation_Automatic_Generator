import json
import re
import sys
import argparse
from docx import Document
import common_functions as cf
import os
import sys



# Method return description about FI flow
def fiMainDescription(Header_Name, Header_Description):

    FI_Descriptoin = ""
    FI_Main_HTML_Data = []
    FI_Main_HTML_Data_Obj = {}
    FI_Main_HTML_Obj = {}

    for i in range(1,4):
        FI_Descriptoin += Header_Name[i] + ": " + Header_Description[i] + "\n"


    FI_Main_HTML_Data_Obj['htmlType'] = 'fiTitle'
    FI_Main_HTML_Data_Obj['txt'] = Header_Name[0] + ": " + Header_Description[0]
    FI_Main_HTML_Data.append(FI_Main_HTML_Data_Obj)
    FI_Main_HTML_Data_Obj = {}
    FI_Main_HTML_Data_Obj['htmlType'] = 'fiStpDsc'
    FI_Main_HTML_Data_Obj['txt'] = FI_Descriptoin
    FI_Main_HTML_Data.append(FI_Main_HTML_Data_Obj)

    FI_Main_HTML_Obj['htmlData'] = FI_Main_HTML_Data



    return FI_Main_HTML_Obj

#############################################################################################
#############################################################################################
###################################                    ######################################
###################################    Main program    ######################################
###################################                    ######################################
#############################################################################################
#############################################################################################




FI_Label = True
FI_Num = 0

FI_Txt_Header = {}
FI_Txt_Header_Exist = False


# Generic titles
FI_Generic_Title_Main = ["Failure No."]
FI_Generic_Title_Des = ["Tested Unit", "Severity", "Name"]
FI_Generic_Step = ["Test ", "DSA"]
FI_Generic_Task = ["LRU", "SAW"]

FI_Generic_Title_Main_Rows = {}
FI_Generic_Title_Des_Rows = {}
FI_Generic_Step_Rows = {}
FI_Generic_Task_Rows = {}


# parser = argparse.ArgumentParser()
# parser.add_argument("file")
# args = parser.parse_args()
# document = Document(args.file)

# cf.init()
#
#
# pathUrl=cf.generate_fi_doc_path(sys.argv[1])
# print ("pathUrl",pathUrl)
# document = Document(pathUrl)
document = Document('C:\\Users\\eden.SPIDERSERVICES\\Desktop\\docx\\DOC-Troubleshooting_table_testing.docx')

FI_Array_List = []
FI_Array = []

tables = document.tables
print(len(tables[0].rows[0].cells))

for table in tables:
    for row in table.rows:
        FI_row = []
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                FI_row.append(paragraph.text)

                if(len(FI_row) == 8):

                    if not FI_Txt_Header_Exist:

                        for i in range(0,len(FI_row)):
                            for title in FI_Generic_Title_Main:
                                if title in FI_row[i]:
                                    FI_Generic_Title_Main_Rows[i] = FI_row[i]

                            for title in FI_Generic_Title_Des:
                                if title in FI_row[i]:
                                    FI_Generic_Title_Des_Rows[i] = FI_row[i]

                            for title in FI_Generic_Step:
                                if title in FI_row[i]:
                                    FI_Generic_Step_Rows[i] = FI_row[i]

                            for title in FI_Generic_Task:
                                if title in FI_row[i]:
                                    FI_Generic_Task_Rows[i] = FI_row[i]

                            FI_Txt_Header[i] = FI_row[i]

                        print(FI_Generic_Title_Main_Rows)
                        print(FI_Generic_Title_Des_Rows)
                        print(FI_Generic_Step_Rows)
                        print(FI_Generic_Task_Rows)

                        FI_Txt_Header_Exist = True

                    else:
                        # Main text info to object (Failure No., Tested Unit, Severity, Platform)
                        FI_Num=0
                        newNumberObj = {}
                        newNumberObj['n'] = str(FI_Num)
                        FI_Num+=1
                        newNumberObj['htmlObj'] = fiMainDescription(FI_Generic_Title_Des_Rows, FI_row)

                        PG_Status = ""
                        for i in range(1, 4):
                            if (FI_row[i] == ""):
                                if (PG_Status == ""):
                                    PG_Status = "title."+str(FI_Txt_Header[i])+"Error"
                                else:
                                    PG_Status += ",title." + str(FI_Txt_Header[i]) + "Error"

                        if (PG_Status == ""):
                            newNumberObj['status'] = "success"
                        else:
                            newNumberObj['status'] = PG_Status

                        FI_Array.append(newNumberObj)
                        Number_Of_Actions = 0

                        for i in range(4, 8):
                            if FI_row[i] != "":
                                Number_Of_Actions +=1

                        for i in range(4, 8):
                            if FI_row[i] != "":
                                newNumberObj = {}
                                newNumberObj['n'] = str(FI_Num)
                                if(i == 4):
                                    # Step
                                    newNumberObj['Y'] = {'to' : str(Number_Of_Actions+1), 'typ' : '0'}
                                    newNumberObj['N'] = {'to' : str(FI_Num+1), 'typ' : '0'}
                                else:
                                    # Task
                                    newNumberObj['Y'] = cf.fiTaskYes2(FI_row[i], str(Number_Of_Actions+1), str(FI_Num+1))
                                    newNumberObj['N'] = { 'typ' : '4' }
                                newNumberObj['htmlObj'] = cf.fiStepDescriptionQuestion(FI_Txt_Header[i] + ": " + FI_row[i])
                                newNumberObj['status'] = "success"
                                FI_Array.append(newNumberObj)
                                FI_Num+=1

                            else:
                                if(i == 4):
                                    # Step is missing
                                    newNumberObj = {}
                                    newNumberObj['n'] = str(FI_Num)
                                    newNumberObj['status'] = "missingStepError"
                                    # FI_Array.append(newNumberObj)
                                    # FI_Num+=1


                        ###### Default end FI numbers ######
                        FI_Array.append({"n": str(FI_Num), "htmlObj": {"htmlData": [{"htmlType": "fiNegEnd"}]},
                                         "N": {"typ": "4"}, "Y": {"typ": "4"}, "status": "success"})
                        FI_Num += 1
                        FI_Array.append({"n": str(FI_Num), "htmlObj": {"htmlData": [{"htmlType": "fiPosEnd"}]},
                                         "N": {"typ": "4"}, "Y": {"typ": "4"}, "status": "success"})
                        FI={"PG":FI_Array,
                            "status":"success"}
                        FI_Array_List.append(FI)
                        FI_Array = []


# os.remove(pathUrl)
# cf.post_api_server(FI_Array_List)
# print(json.dumps(FI_Array_List, indent=4, sort_keys=True))