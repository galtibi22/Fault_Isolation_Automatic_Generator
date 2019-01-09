import json
import sys
import argparse
from docx import Document

# Method check if string has Yes/No
def checkYesNo(str):
    if (str == "Yes" or str == "yes" or str == "No" or str == "no"):
        return True
    return False


# Method return to Step description and question from string
def fiStepDescriptionQuestion(str):
    htmlDataObj = {}
    htmlDataObjAraay = []

    arrDesQue = str.split('.')

    if('?' not in str):
        objDes = {'htmlType':'fiStpDsc', 'txt' : str}
        htmlDataObjAraay.append(objDes)
    else:
        newDes = ""
        for description in arrDesQue[:-1]:
            newDes += description+"."
        objDes = {'htmlType':'fiStpDsc', 'txt' : newDes}
        htmlDataObjAraay.append(objDes)
        objQue = {'fiStpQst':'fiStpDsc', 'txt' : arrDesQue[-1]}
        htmlDataObjAraay.append(objQue)

    htmlDataObj['htmlData'] = htmlDataObjAraay
    return htmlDataObj



# Method return to Task description and question from string
def fiTaskDescriptionQuestion(str):
    htmlDataObj = {}
    htmlDataObjAraay = []

    arrDesQue = str.split('.')

    if('?' not in str):
        objDes = {'htmlType':'fiStpDsc', 'txt' : str}
        htmlDataObjAraay.append(objDes)
    else:
        newDes = ""
        for description in arrDesQue[:-1]:
            newDes += description+"."
        objDes = {'htmlType':'fiStpDsc', 'txt' : newDes}
        htmlDataObjAraay.append(objDes)

    htmlDataObj['htmlData'] = htmlDataObjAraay
    return htmlDataObj


# Method return remove characters
def removeCharactersFunc(str):
    for rchar in removeCharacters:
        str = str.replace(rchar,'')
    return str


# Method return Task Object to 'Y' (Yes option only)
def fiTaskYes(str, yesOption, noOption):
    yesObj = {}
    arrDesQue = str.split('.')
    newDes = ""

    if('?' not in str):
        newDes = str
    else:
        yesObj['msgRtIx'] = arrDesQue[-1]
        for description in arrDesQue[:-1]:
            newDes += description+"."

    yesObj['to'] = removeCharactersFunc(newDes[newDes.find("(")+1:newDes.find(")")])
    yesObj['rtN'] = removeCharactersFunc(noOption)
    yesObj['msgRt'] = "1"
    yesObj['typ'] = "1"
    yesObj['tskNm'] = newDes[:newDes.find("(")]
    yesObj['rtY'] = removeCharactersFunc(removeCharactersFunc(yesOption))

    return yesObj

    return {'to' : str, 'typ' : '0'}

############# Main program ###############

removeCharacters = ['\n','\t','\u200e']
IsFaultIsolation = False
FI_Dict = {}
FI_Label = True
FI_Descriptoin = ""
FI_Num = 0

parser = argparse.ArgumentParser()
parser.add_argument("file")
args = parser.parse_args()
document = Document(args.file)

# document = Document('DOC-Left-Wagon.docx')


FI_Array = []

# Main text info to object
FI_Main_Info = {}
FI_Main_Info['n'] = str(FI_Num)
FI_Num+=1
FI_Main_HTML_Obj ={}
FI_Main_HTML_Data = []
FI_Main_HTML_Data_Obj = {}

for para in document.paragraphs:
    if (FI_Label == True):
        FI_Main_HTML_Data_Obj['htmlType'] = 'fiTitle'
        FI_Main_HTML_Data_Obj['txt'] = para.text
        FI_Label = False
        FI_Main_HTML_Data.append(FI_Main_HTML_Data_Obj)
        FI_Main_HTML_Data_Obj = {}
    else:
        FI_Descriptoin+=para.text + "\n"

FI_Main_HTML_Data_Obj['htmlType'] = 'fiDes'
FI_Main_HTML_Data_Obj['txt'] = FI_Descriptoin
FI_Main_HTML_Data.append(FI_Main_HTML_Data_Obj)

FI_Main_HTML_Obj['htmlData'] = FI_Main_HTML_Data
FI_Main_Info['htmlObj'] = FI_Main_HTML_Obj
FI_Array.append(FI_Main_Info)

tables = document.tables
for table in tables:
    for row in table.rows:
        FI_row = []
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if not checkYesNo(paragraph.text):
                    # print(paragraph.text)
                    FI_row.append(paragraph.text)
                    if len(FI_row) == 3:
                        newNumberObj = {}
                        newNumberObj['n'] = str(FI_Num)
                        if(('(' not in  FI_row[0]) and (')' not in  FI_row[0])):
                            # Step
                            newNumberObj['Y'] = {'to' : removeCharactersFunc(FI_row[1]), 'typ' : '0'}
                            newNumberObj['N'] = {'to' : removeCharactersFunc(FI_row[2]), 'typ' : '0'}
                            newNumberObj['htmlObj'] = fiStepDescriptionQuestion(FI_row[0])
                        else:
                            # Task
                            newNumberObj['Y'] = fiTaskYes(FI_row[0], FI_row[1], FI_row[2])
                            newNumberObj['N'] = { 'typ' : '4' }
                            newNumberObj['htmlObj'] = fiTaskDescriptionQuestion(FI_row[0])
                        FI_Array.append(newNumberObj)
                        FI_Num+=1


###### Default end FI numbers ######
FI_Array.append({ "n": str(FI_Num), "htmlObj": { "htmlData": [ { "htmlType": "fiNegEnd" } ] }, "N": { "typ": "4" }, "Y": { "typ": "4" } })
FI_Num+=1
FI_Array.append({ "n": str(FI_Num), "htmlObj": { "htmlData": [ { "htmlType": "fiPosEnd" } ] }, "N": { "typ": "4" }, "Y": { "typ": "4" } })

print(json.dumps(FI_Array, indent=4, sort_keys=True))