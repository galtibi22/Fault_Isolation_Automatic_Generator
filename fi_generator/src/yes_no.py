import json
import argparse
from docx import Document
import common_functions as cf
import os
import sys
#############################################################################################
#############################################################################################
###################################                    ######################################
###################################    Main program    ######################################
###################################                    ######################################
#############################################################################################
#############################################################################################

FI_Label = True
FI_Num = 0


path=cf.init()
pathUrl=cf.generate_fi_doc_path(sys.argv[1])
document = Document(pathUrl)


FI_Array = []

# Main text info to object
FI_Descriptoin = ""
FI_Main_Info = {}
FI_Main_Info['n'] = str(FI_Num)
FI_Num+=1
FI_Main_HTML_Obj ={}
FI_Main_HTML_Data = []
FI_Main_HTML_Data_Obj = {}
FI_Save_Errors = ""

for para in document.paragraphs:
    if (FI_Label == True):
        FI_Main_HTML_Data_Obj['htmlType'] = 'fiTitle'
        FI_Main_HTML_Data_Obj['txt'] = para.text
        if FI_Main_HTML_Data_Obj['txt'] == "":
            FI_Save_Errors = "titleIsMissing"
        FI_Label = False
        FI_Main_HTML_Data.append(FI_Main_HTML_Data_Obj)
        FI_Main_HTML_Data_Obj = {}
    else:
        FI_Descriptoin+=para.text + "\n"

FI_Main_HTML_Data_Obj['htmlType'] = 'fiStpDsc'
FI_Main_HTML_Data_Obj['txt'] = FI_Descriptoin
FI_Main_HTML_Data.append(FI_Main_HTML_Data_Obj)

FI_Main_HTML_Obj['htmlData'] = FI_Main_HTML_Data
FI_Main_Info['htmlObj'] = FI_Main_HTML_Obj
if (FI_Save_Errors==""):
    FI_Main_Info['status'] = "success"
else:
    FI_Main_Info['status'] = FI_Save_Errors

if (FI_Main_Info['htmlObj']):
    print("eden")

FI_Array.append(FI_Main_Info)

tables = document.tables

# print(len(tables[0].rows[0].cells))
for table in tables:
    for row in table.rows:
        FI_row = []
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if not cf.checkYesNo(paragraph.text):
                    # print(paragraph.text)
                    FI_row.append(paragraph.text)
                    if len(FI_row) == 3:
                        des=FI_row[0]
                        newNumberObj = {}
                        newNumberObj['n'] = str(FI_Num)
                        if(('(' not in  FI_row[0]) and (')' not in  FI_row[0])):
                            # Test
                            newNumberObj['Y'] = {'to' : cf.removeCharsFunc(FI_row[1]), 'typ' : '0'}
                            newNumberObj['N'] = {'to' : cf.removeCharsFunc(FI_row[2]), 'typ' : '0'}
                            newNumberObj['type'] = "test"
                            newNumberObj['htmlObj'] = cf.fiTestDescriptionQuestion(FI_row[0])
                        else:
                            # Task
                            newNumberObj['Y'] = cf.fiTaskYes(FI_row[0], FI_row[1], FI_row[2])
                            newNumberObj['N'] = { 'typ' : '4' }
                            newNumberObj['htmlObj'] = cf.fiTaskDescriptionQuestion(FI_row[0])
                            newNumberObj['type'] = "task"
                        newNumberObj['status'] = "success"
                        FI_Array.append(newNumberObj)
                        FI_Num+=1


###### Default end FI numbers ######
FI_Array.append({ "n": str(FI_Num), "htmlObj": { "htmlData": [ { "htmlType": "fiNegEnd" } ] }, "N": { "typ": "4" }, "Y": { "typ": "4" }, "status": "success" })
FI_Num+=1
FI_Array.append({ "n": str(FI_Num), "htmlObj": { "htmlData": [ { "htmlType": "fiPosEnd" } ] }, "N": { "typ": "4" }, "Y": { "typ": "4" }, "status": "success" })

FI_Array = cf.replaceLineToNumber(FI_Array)
FI_Array = cf.recheckFlow(FI_Array)

fis=[{
    'PG':FI_Array,
    "status":"success"
}]

os.remove(pathUrl)
cf.post_api_server(fis)
# print(json.dumps(fis, indent=4, sort_keys=True))

